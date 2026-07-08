"""
Offline bulk ingestion for the HSE Assistant.
Run this ON YOUR MACHINE (not Streamlit Cloud) to load all SharePoint docs into the
SAME cloud Qdrant the app reads. Resumable: safe to stop and re-run.

  pip install requests qdrant-client pdfplumber python-docx tqdm
  # set the env vars below (or edit the CONFIG block), then:
  python ingest.py

Uses mistral-embed (1024-dim) so vectors match the app's query embeddings.
"""
import os, io, re, json, time, uuid
from pathlib import Path
from urllib.parse import urlparse, unquote

import requests
import pdfplumber
import docx as _docx
from tqdm import tqdm
from qdrant_client import QdrantClient
from qdrant_client.models import Distance, VectorParams, PointStruct

# ─────────────────────────────── CONFIG (env vars, or edit the defaults) ───────
MISTRAL_KEY   = os.getenv("MISTRAL_API_KEY",  "<your-mistral-key>")
QDRANT_URL    = os.getenv("QDRANT_URL",        "https://XXXX.sa-east-1-0.aws.cloud.qdrant.io:6333")
QDRANT_KEY    = os.getenv("QDRANT_API_KEY",    "<your-qdrant-key>")
TENANT_ID     = os.getenv("TENANT_ID",         "<tenant-id>")
CLIENT_ID     = os.getenv("CLIENT_ID",         "<client-id>")
CLIENT_SECRET = os.getenv("CLIENT_SECRET",     "<client-secret>")
SITE_URL      = os.getenv("SITE_URL",          "https://vestas.sharepoint.com/sites/GlobalQHSE-hub/HSEN Legacy")

COLLECTION    = "hse_docs"
EMBED_MODEL   = "mistral-embed"        # 1024-dim (must match the app)
CHUNK_SIZE    = 1000
CHUNK_OVERLAP = 200
EMBED_BATCH   = 16
SUPPORTED     = {".pdf", ".docx", ".txt", ".md"}
GRAPH         = "https://graph.microsoft.com/v1.0"
STATE_FILE    = "processed.json"       # remembers finished item_ids for resume

# ─────────────────────────────── Qdrant ────────────────────────────────────────
qc = QdrantClient(url=QDRANT_URL, api_key=QDRANT_KEY, timeout=120, check_compatibility=False)

def ensure_collection():
    names = [c.name for c in qc.get_collections().collections]
    if COLLECTION not in names:
        qc.create_collection(COLLECTION,
                             vectors_config=VectorParams(size=1024, distance=Distance.COSINE))
        print(f'Created collection "{COLLECTION}"')

# ─────────────────────────────── resume state ──────────────────────────────────
def load_state():
    if os.path.exists(STATE_FILE):
        return set(json.load(open(STATE_FILE)))
    return set()

def save_state(done):
    json.dump(sorted(done), open(STATE_FILE, "w"))

# ─────────────────────────────── Mistral embed (429-safe) ──────────────────────
def embed(texts):
    out = []
    for i in range(0, len(texts), EMBED_BATCH):
        batch = texts[i:i + EMBED_BATCH]
        for attempt in range(8):
            r = requests.post("https://api.mistral.ai/v1/embeddings",
                              headers={"Authorization": f"Bearer {MISTRAL_KEY}"},
                              json={"model": EMBED_MODEL, "input": batch}, timeout=120)
            if r.status_code == 429 or r.status_code >= 500:
                wait = float(r.headers.get("Retry-After", 2 ** attempt))
                time.sleep(min(wait, 60)); continue
            r.raise_for_status()
            out.extend(d["embedding"] for d in r.json()["data"]); break
        else:
            raise RuntimeError("Mistral rate-limited after retries")
        time.sleep(0.2)
    return out

# ─────────────────────────────── SharePoint (app-only) ─────────────────────────
def graph_token():
    r = requests.post(f"https://login.microsoftonline.com/{TENANT_ID}/oauth2/v2.0/token",
                      data={"client_id": CLIENT_ID, "client_secret": CLIENT_SECRET,
                            "grant_type": "client_credentials",
                            "scope": "https://graph.microsoft.com/.default"}, timeout=30)
    r.raise_for_status()
    return r.json()["access_token"]

def gget(url, token):
    for a in range(6):
        r = requests.get(url, headers={"Authorization": f"Bearer {token}"}, timeout=90)
        if r.status_code in (429, 503, 504):
            time.sleep(int(r.headers.get("Retry-After", 2 ** a))); continue
        return r
    return r

def parse_site_url(url):
    m = re.match(r"https?://([^/]+)((?:/sites/[^/?#]+)?)", url, re.I)
    host, site = (m.group(1), m.group(2) or "/") if m else (None, None)
    path = unquote(urlparse(url).path)
    path = re.sub(r"/Forms/[^/]*$", "", path)
    path = re.sub(r"/[^/]+\.aspx$", "", path).rstrip("/")
    return host, site, path

def resolve_drive(site_path, folder_path, drives, default_drive):
    rest = folder_path[len(site_path):].strip("/") if folder_path.startswith(site_path) else ""
    segs = [s for s in rest.split("/") if s]
    if segs:
        for d in drives:
            if d.get("name", "").lower() == segs[0].lower():
                return d, "/".join(segs[1:])
    for d in drives:
        dp = unquote(urlparse(d.get("webUrl", "")).path).rstrip("/")
        if dp and folder_path.startswith(dp):
            return d, folder_path[len(dp):].strip("/")
    return default_drive, rest

def iter_files(site_id, drive_id, token, folder="root"):
    stack = [folder]
    while stack:
        fid = stack.pop()
        url = (f"{GRAPH}/sites/{site_id}/drives/{drive_id}/items/{fid}/children"
               "?$top=200&$select=id,name,file,folder,webUrl,size")
        while url:
            r = gget(url, token); r.raise_for_status(); data = r.json()
            for it in data.get("value", []):
                if it.get("folder"):
                    stack.append(it["id"])
                elif it.get("file"):
                    yield it
            url = data.get("@odata.nextLink")

def download(drive_id, item_id, token):
    r = gget(f"{GRAPH}/drives/{drive_id}/items/{item_id}/content", token)
    r.raise_for_status()
    return r.content

# ─────────────────────────────── parse + chunk ─────────────────────────────────
def extract_text(name, data):
    ext = Path(name).suffix.lower()
    try:
        if ext == ".pdf":
            parts = []
            with pdfplumber.open(io.BytesIO(data)) as pdf:
                for page in pdf.pages:
                    t = page.extract_text() or ""
                    if t.strip():
                        parts.append(t.strip())
                    try:
                        for tbl in page.extract_tables():
                            rows = [" | ".join((c or "").strip() for c in row) for row in tbl]
                            rows = [r for r in rows if r.strip(" |")]
                            if rows:
                                parts.append("[TABLE]\n" + "\n".join(rows))
                    except Exception:
                        pass
            return "\n\n".join(parts)
        if ext == ".docx":
            d = _docx.Document(io.BytesIO(data))
            paras = [p.text for p in d.paragraphs if p.text.strip()]
            for tbl in d.tables:
                for row in tbl.rows:
                    cells = [c.text.strip() for c in row.cells]
                    if any(cells):
                        paras.append(" | ".join(cells))
            return "\n".join(paras)
        if ext in (".txt", ".md"):
            return data.decode("utf-8", errors="ignore")
    except Exception as e:
        print(" parse error", name, e)
    return ""

def chunk(text):
    text = re.sub(r"\s+\n", "\n", text).strip()
    out, i = [], 0
    while i < len(text):
        piece = text[i:i + CHUNK_SIZE]
        if len(piece.strip()) > 50:
            out.append(piece)
        i += CHUNK_SIZE - CHUNK_OVERLAP
    return out

def point_id(item_id, ordinal):
    return str(uuid.uuid5(uuid.NAMESPACE_URL, f"{item_id}:{ordinal}"))

# ─────────────────────────────── main ──────────────────────────────────────────
def main():
    ensure_collection()
    done = load_state()
    token = graph_token()

    host, site, folder_path = parse_site_url(SITE_URL)
    site_id = gget(f"{GRAPH}/sites/{host}:{site}", token).json()["id"]
    drives = gget(f"{GRAPH}/sites/{site_id}/drives", token).json().get("value", [])
    default_drive = gget(f"{GRAPH}/sites/{site_id}/drive", token).json()
    drive, subfolder = resolve_drive(site, folder_path, drives, default_drive)
    drive_id = drive["id"]
    print(f'Library: {drive.get("name")}  subfolder: "{subfolder}"')

    start = "root"
    if subfolder:
        r = gget(f"{GRAPH}/sites/{site_id}/drives/{drive_id}/root:/{subfolder}", token)
        if r.status_code == 200:
            start = r.json()["id"]

    print("Listing files…")
    files = [it for it in iter_files(site_id, drive_id, token, start)
             if Path(it["name"]).suffix.lower() in SUPPORTED]
    todo = [it for it in files if it["id"] not in done]
    print(f"{len(files)} files total · {len(done)} already done · {len(todo)} to do")

    indexed = notext = failed = chunks_total = 0
    for it in tqdm(todo, desc="Indexing", unit="doc"):
        try:
            text = extract_text(it["name"], download(drive_id, it["id"], token))
            pieces = chunk(text)
            if not pieces:
                notext += 1
            else:
                vecs = embed(pieces)
                pts = [PointStruct(id=point_id(it["id"], j), vector=v,
                                   payload={"item_id": it["id"], "name": it["name"],
                                            "web_url": it.get("webUrl", ""), "text": ch})
                       for j, (ch, v) in enumerate(zip(pieces, vecs))]
                qc.upsert(COLLECTION, points=pts)
                indexed += 1; chunks_total += len(pts)
            done.add(it["id"])
            if len(done) % 25 == 0:            # checkpoint often so resume is precise
                save_state(done)
        except Exception as e:
            failed += 1
            tqdm.write(f"FAILED {it['name']}: {e}")
    save_state(done)
    print(f"\nDone. {indexed} indexed · {notext} no-text · {failed} failed · "
          f"{chunks_total} new chunks. Collection now has {qc.count(COLLECTION).count} points.")

if __name__ == "__main__":
    main()
