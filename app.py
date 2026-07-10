"""
HSE Assistant — shared team app (rich UI + branding + memory-safe batch sync)
────────────────────────────────────────────────────────────────────────────
Public docs → one shared index → give the URL to your team, they just ask.

Bulk loading 1900+ docs: prefer ingest.py (locally or GitHub Actions).
In-app Sync is memory-safe: it indexes a small slice per click and stops.
Resume is reliable: it fetches the set of already-indexed doc IDs once and
skips them, so clicking again advances to NEW docs (never restarts from doc 1).
"""
import os, io, re, json, time, uuid, base64
from pathlib import Path
from urllib.parse import urlparse, urlunparse, unquote

import requests
import streamlit as st
import pdfplumber
import docx as _docx
from qdrant_client import QdrantClient
from qdrant_client.models import (Distance, VectorParams, PointStruct,
                                  Filter, FieldCondition, MatchValue)

st.set_page_config(page_title="HSE Assistant", page_icon="🦺", layout="wide")

# ─────────────────────────────────────────────────────────── config (st.secrets)
def cfg(key, default=None):
    if key in st.secrets:
        return st.secrets[key]
    return os.getenv(key, default)

def normalize_qdrant_url(u):
    u = (u or "").strip().rstrip("/")
    if not u:
        return u
    if not re.match(r"^https?://", u, re.I):
        u = "https://" + u
    p = urlparse(u)
    netloc = p.netloc if ":" in p.netloc else p.netloc + ":6333"
    return urlunparse((p.scheme, netloc, "", "", "", ""))

MISTRAL_KEY   = cfg("MISTRAL_API_KEY", "")
QDRANT_URL    = normalize_qdrant_url(cfg("QDRANT_URL", ""))
QDRANT_KEY    = cfg("QDRANT_API_KEY", "")
ADMIN_PW      = cfg("ADMIN_PASSWORD", "")

TENANT_ID     = cfg("TENANT_ID", "")
CLIENT_ID     = cfg("CLIENT_ID", "")
CLIENT_SECRET = cfg("CLIENT_SECRET", "")
SITE_URL      = cfg("SITE_URL", "")

CHAT_MODEL    = cfg("CHAT_MODEL", "open-mistral-7b")
EMBED_MODEL   = "mistral-embed"
COLLECTION    = "hse_docs"
CONFIG_COLL   = "app_config"
TOP_K         = int(cfg("TOP_K", 6))
CHUNK_SIZE    = 1000
CHUNK_OVERLAP = 200
EMBED_BATCH   = 16
SUPPORTED     = {".pdf", ".docx", ".txt", ".md"}
GRAPH         = "https://graph.microsoft.com/v1.0"

DEFAULT_BRAND = {
    "title":    cfg("APP_TITLE", "HSE Assistant"),
    "subtitle": cfg("APP_SUBTITLE", "Ask anything about your HSE documents"),
    "accent":   cfg("ACCENT_COLOR", "#0F9D8C"),
    "logo_url": cfg("LOGO_URL", ""),
    "logo_b64": "",
    "logo_mime": "",
}

SUGGESTED = [
    "What PPE is required for confined space entry?",
    "Summarise the permit-to-work procedure.",
    "What are the steps in incident reporting?",
    "What does the standard say about working at height?",
]

# ─────────────────────────────────────────────────────────── Qdrant
@st.cache_resource(show_spinner=False)
def qdrant():
    return QdrantClient(url=QDRANT_URL, api_key=QDRANT_KEY,
                        timeout=60, check_compatibility=False)

def ensure_collection():
    names = [x.name for x in qdrant().get_collections().collections]
    if COLLECTION not in names:
        qdrant().create_collection(COLLECTION,
                                   vectors_config=VectorParams(size=1024, distance=Distance.COSINE))
    # payload index on item_id makes filtering/skip fast (idempotent)
    try:
        qdrant().create_payload_index(collection_name=COLLECTION,
                                      field_name="item_id", field_schema="keyword")
    except Exception:
        pass

def index_count():
    try:
        return qdrant().count(COLLECTION).count
    except Exception:
        return 0

def indexed_item_ids():
    """Return the SET of SharePoint item_ids already present in the index.
    Fetched in one pass (paged) so batch-sync can skip them reliably & fast."""
    ids = set()
    try:
        offset = None
        while True:
            points, offset = qdrant().scroll(
                collection_name=COLLECTION, limit=1000,
                with_payload=["item_id"], with_vectors=False, offset=offset)
            for p in points:
                iid = (p.payload or {}).get("item_id")
                if iid:
                    ids.add(iid)
            if offset is None:
                break
    except Exception:
        pass
    return ids

# ── branding stored in Qdrant so it persists + shows to everyone ──
def _ensure_config_coll():
    names = [x.name for x in qdrant().get_collections().collections]
    if CONFIG_COLL not in names:
        qdrant().create_collection(CONFIG_COLL,
                                   vectors_config=VectorParams(size=1, distance=Distance.DOT))

@st.cache_data(ttl=120, show_spinner=False)
def load_brand():
    brand = dict(DEFAULT_BRAND)
    try:
        _ensure_config_coll()
        res = qdrant().retrieve(CONFIG_COLL, ids=[1], with_payload=True)
        if res and res[0].payload:
            brand.update({k: v for k, v in res[0].payload.items() if v})
    except Exception:
        pass
    return brand

def save_brand(brand):
    _ensure_config_coll()
    qdrant().upsert(CONFIG_COLL, points=[PointStruct(id=1, vector=[1.0], payload=brand)])
    load_brand.clear()

def logo_src(brand):
    if brand.get("logo_b64"):
        return f"data:{brand.get('logo_mime','image/png')};base64,{brand['logo_b64']}"
    if brand.get("logo_url"):
        return brand["logo_url"]
    if os.path.exists("assets/logo.png"):
        b64 = base64.b64encode(open("assets/logo.png", "rb").read()).decode()
        return f"data:image/png;base64,{b64}"
    return ""

# ─────────────────────────────────────────────────────────── Mistral helpers
def mistral_embed(texts):
    if isinstance(texts, str):
        texts = [texts]
    out = []
    for i in range(0, len(texts), EMBED_BATCH):
        batch = texts[i:i + EMBED_BATCH]
        for attempt in range(6):
            r = requests.post("https://api.mistral.ai/v1/embeddings",
                              headers={"Authorization": f"Bearer {MISTRAL_KEY}"},
                              json={"model": EMBED_MODEL, "input": batch}, timeout=120)
            if r.status_code == 429 or r.status_code >= 500:
                time.sleep(min(float(r.headers.get("Retry-After", 2 ** attempt)), 30)); continue
            r.raise_for_status()
            out.extend([d["embedding"] for d in r.json()["data"]]); break
        else:
            raise RuntimeError("Mistral embeddings rate-limited after retries")
        time.sleep(0.3)
    return out

def mistral_chat(system, user, temperature=0.1):
    r = requests.post("https://api.mistral.ai/v1/chat/completions",
                      headers={"Authorization": f"Bearer {MISTRAL_KEY}"},
                      json={"model": CHAT_MODEL, "temperature": temperature, "max_tokens": 1200,
                            "messages": [{"role": "system", "content": system},
                                         {"role": "user", "content": user}]}, timeout=180)
    r.raise_for_status()
    return r.json()["choices"][0]["message"]["content"]

# ─────────────────────────────────────────────────────────── SharePoint (app-only)
def graph_app_token():
    r = requests.post(f"https://login.microsoftonline.com/{TENANT_ID}/oauth2/v2.0/token",
                      data={"client_id": CLIENT_ID, "client_secret": CLIENT_SECRET,
                            "grant_type": "client_credentials",
                            "scope": "https://graph.microsoft.com/.default"}, timeout=30)
    if r.status_code != 200:
        raise RuntimeError(f"Graph token failed: {r.status_code} {r.text[:200]}")
    return r.json()["access_token"]

def gget(url, token):
    for a in range(5):
        r = requests.get(url, headers={"Authorization": f"Bearer {token}"}, timeout=60)
        if r.status_code in (429, 503, 504):
            time.sleep(int(r.headers.get("Retry-After", 2 ** a))); continue
        return r
    return r

def parse_site_url(url):
    m = re.match(r"https?://([^/]+)((?:/sites/[^/?#]+)?)", url, re.I)
    host, site = (m.group(1), m.group(2) or "/") if m else (None, None)
    path = unquote(urlparse(url).path)
    path = re.sub(r"/Forms/[^/]*$", "", path)
    path = re.sub(r"/[^/]+\.aspx$", "", path).rstrip("/")
    return host, site, path

def resolve_drive(site_path, folder_path, drives, default_drive, log):
    rest = folder_path[len(site_path):].strip("/") if folder_path.startswith(site_path) else ""
    segments = [s for s in rest.split("/") if s]
    if segments:
        for d in drives:
            if d.get("name", "").lower() == segments[0].lower():
                log(f"matched library by name: {d.get('name')}"); return d, "/".join(segments[1:])
    for d in drives:
        dp = unquote(urlparse(d.get("webUrl", "")).path).rstrip("/")
        if dp and folder_path.startswith(dp):
            log(f"matched library by url: {d.get('name')}"); return d, folder_path[len(dp):].strip("/")
    log(f"no library match; using default '{default_drive.get('name')}' + subfolder '{rest}'")
    return default_drive, rest

def iter_files(site_id, drive_id, token, folder="root"):
    stack = [folder]
    while stack:
        fid = stack.pop()
        url = (f"{GRAPH}/sites/{site_id}/drives/{drive_id}/items/{fid}/children"
               "?$top=200&$select=id,name,file,folder,webUrl,size")
        while url:
            r = gget(url, token); r.raise_for_status(); data = r.json()
            for it in data.get("value", []):
                if it.get("folder"):
                    stack.append(it["id"])
                elif it.get("file"):
                    yield it
            url = data.get("@odata.nextLink")

def download(drive_id, item_id, token):
    r = gget(f"{GRAPH}/drives/{drive_id}/items/{item_id}/content", token)
    r.raise_for_status()
    return r.content

# ─────────────────────────────────────────────────────────── parsing + chunking
def extract_text(name, data):
    ext = Path(name).suffix.lower()
    try:
        if ext == ".pdf":
            parts = []
            with pdfplumber.open(io.BytesIO(data)) as pdf:
                for page in pdf.pages:
                    t = page.extract_text() or ""
                    if t.strip():
                        parts.append(t.strip())
                    try:
                        for tbl in page.extract_tables():
                            rows = [" | ".join((c or "").strip() for c in row) for row in tbl]
                            rows = [r for r in rows if r.strip(" |")]
                            if rows:
                                parts.append("[TABLE]\n" + "\n".join(rows))
                    except Exception:
                        pass
            return "\n\n".join(parts)
        if ext == ".docx":
            d = _docx.Document(io.BytesIO(data))
            paras = [p.text for p in d.paragraphs if p.text.strip()]
            for tbl in d.tables:
                for row in tbl.rows:
                    cells = [c.text.strip() for c in row.cells]
                    if any(cells):
                        paras.append(" | ".join(cells))
            return "\n".join(paras)
        if ext in (".txt", ".md"):
            return data.decode("utf-8", errors="ignore")
    except Exception as e:
        print("parse error", name, e)
    return ""

def chunk(text):
    text = re.sub(r"\s+\n", "\n", text).strip()
    out, i = [], 0
    while i < len(text):
        piece = text[i:i + CHUNK_SIZE]
        if len(piece.strip()) > 50:
            out.append(piece)
        i += CHUNK_SIZE - CHUNK_OVERLAP
    return out

def point_id(item_id, ordinal):
    return str(uuid.uuid5(uuid.NAMESPACE_URL, f"{item_id}:{ordinal}"))

# ─────────────────────────────────────────────────────────── ingestion (batch/slice)
def _resolve_target(token, log):
    host, site, folder_path = parse_site_url(SITE_URL)
    if not host:
        raise RuntimeError(f"Could not parse SITE_URL: {SITE_URL}")
    site_resp = gget(f"{GRAPH}/sites/{host}:{site}", token).json()
    if "error" in site_resp:
        raise RuntimeError(f"Site access failed ({site}): "
                           f"{site_resp['error'].get('message', site_resp['error'])}. "
                           "If 403/accessDenied, the app isn't granted on this site.")
    site_id = site_resp["id"]
    drives = gget(f"{GRAPH}/sites/{site_id}/drives", token).json().get("value", [])
    if not drives:
        raise RuntimeError("No document libraries found (or no access).")
    default_drive = gget(f"{GRAPH}/sites/{site_id}/drive", token).json()
    log("Libraries visible: " + ", ".join(d.get("name", "?") for d in drives))
    drive, subfolder = resolve_drive(site, folder_path, drives, default_drive, log)
    drive_id = drive["id"]
    start_folder = "root"
    if subfolder:
        r = gget(f"{GRAPH}/sites/{site_id}/drives/{drive_id}/root:/{subfolder}", token)
        if r.status_code == 200:
            start_folder = r.json()["id"]
        else:
            log(f"subfolder '{subfolder}' not found — crawling library root")
    return site_id, drive_id, start_folder

def index_one(drive_id, it, token, log):
    text = extract_text(it["name"], download(drive_id, it["id"], token))
    pieces = chunk(text)
    if not pieces:
        log(f"skip (no text): {it['name']}")
        return 0
    vecs = mistral_embed(pieces)
    pts = [PointStruct(id=point_id(it["id"], j), vector=v,
                       payload={"item_id": it["id"], "name": it["name"],
                                "web_url": it.get("webUrl", ""), "text": ch})
           for j, (ch, v) in enumerate(zip(pieces, vecs))]
    qdrant().upsert(COLLECTION, points=pts)
    log(f"indexed: {it['name']} ({len(pts)} chunks)")
    return len(pts)

def sync_batch(limit, progress, log, force=False):
    """Index up to `limit` NOT-yet-indexed files, then STOP. Returns files still remaining."""
    ensure_collection()
    token = graph_app_token()
    site_id, drive_id, start_folder = _resolve_target(token, log)

    log("Listing files…")
    files = [it for it in iter_files(site_id, drive_id, token, start_folder)
             if Path(it["name"]).suffix.lower() in SUPPORTED]
    log(f"{len(files)} supported files found. Checking what's already indexed…")

    done_ids = set() if force else indexed_item_ids()
    todo = [it for it in files if it["id"] not in done_ids]
    log(f"{len(done_ids)} docs already indexed · {len(todo)} still to do.")

    this_batch = todo[:limit]
    indexed = notext = failed = chunks_total = 0
    for k, it in enumerate(this_batch, 1):
        try:
            c = index_one(drive_id, it, token, log)
            if c:
                indexed += 1; chunks_total += c
            else:
                notext += 1
        except Exception as e:
            failed += 1; log(f"FAILED {it['name']}: {e}")
        progress(k, len(this_batch), indexed, notext, failed, chunks_total)

    remaining = max(len(todo) - len(this_batch), 0)
    log(f"Batch done. {indexed} indexed · {notext} no-text · {failed} failed · "
        f"{chunks_total} chunks. ~{remaining} files remaining.")
    return remaining

# ─────────────────────────────────────────────────────────── RAG
def answer(question):
    if any(k in question.lower() for k in ("how many document", "number of document")):
        return f"There are **{index_count():,} chunks** indexed in the knowledge base.", []
    qv = mistral_embed(question)[0]
    hits = qdrant().query_points(COLLECTION, query=qv, limit=TOP_K, with_payload=True).points
    if not hits:
        return "I don't have any indexed documents that cover this yet.", []
    ctx, refs, seen = [], [], set()
    for i, h in enumerate(hits):
        p = h.payload
        ctx.append(f"[{i+1}] (from: {p.get('name', '?')})\n{p.get('text', '')}")
        key = (p.get("name"), p.get("web_url", ""))
        if key not in seen:
            seen.add(key); refs.append({"name": p.get("name", "?"), "web_url": p.get("web_url", "")})
    system = ("You are an HSE (Health, Safety & Environment) assistant. Answer ONLY from the "
              "context. Cite the source filename for each fact. If the answer isn't in the "
              "context, say so — never invent information.")
    user = "Context:\n" + "\n\n---\n\n".join(ctx) + f"\n\nQuestion: {question}\n\nAnswer:"
    return mistral_chat(system, user), refs

# ═══════════════════════════════════════════════════════════ UI
brand = load_brand()
ACCENT = brand.get("accent", "#0F9D8C")

st.markdown(f"""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap');
:root {{ --accent:{ACCENT}; }}
html, body, [class*="css"] {{ font-family:'Inter',sans-serif; }}
#MainMenu, footer {{ visibility:hidden; }}
.block-container {{ padding-top:1.4rem; max-width:1100px; }}
.app-header {{
  display:flex; align-items:center; gap:16px; padding:16px 20px; margin-bottom:6px;
  background:linear-gradient(135deg, color-mix(in srgb, var(--accent) 14%, transparent), transparent);
  border:1px solid color-mix(in srgb, var(--accent) 25%, transparent); border-radius:16px;
}}
.app-logo {{ height:46px; width:auto; border-radius:10px; object-fit:contain; }}
.app-logo-fallback {{ height:46px; width:46px; border-radius:12px; display:flex; align-items:center;
  justify-content:center; font-size:24px; background:var(--accent); color:#fff; }}
.app-title {{ font-size:22px; font-weight:800; line-height:1.1; }}
.app-sub {{ font-size:13px; opacity:.7; margin-top:2px; }}
.app-badge {{ margin-left:auto; font-size:12px; font-weight:600; padding:6px 12px; border-radius:20px;
  background:color-mix(in srgb, var(--accent) 16%, transparent);
  color:var(--accent); border:1px solid color-mix(in srgb, var(--accent) 30%, transparent); white-space:nowrap; }}
div[data-testid="stChatMessage"] {{ border-radius:14px; padding:4px 6px; }}
.stButton>button {{ border-radius:12px; border:1px solid color-mix(in srgb, var(--accent) 35%, transparent);
  font-weight:600; transition:all .15s ease; }}
.stButton>button:hover {{ border-color:var(--accent); color:var(--accent); transform:translateY(-1px); }}
.suggest-label {{ font-size:12px; font-weight:700; letter-spacing:.08em; text-transform:uppercase;
  opacity:.55; margin:6px 0 4px; }}
.stTabs [data-baseweb="tab-list"] {{ gap:4px; }}
.stTabs [data-baseweb="tab"] {{ border-radius:10px 10px 0 0; font-weight:600; }}
</style>
""", unsafe_allow_html=True)

src = logo_src(brand)
logo_html = (f'<img class="app-logo" src="{src}"/>' if src
             else '<div class="app-logo-fallback">🦺</div>')
n_chunks = index_count()
st.markdown(f"""
<div class="app-header">
  {logo_html}
  <div>
    <div class="app-title">{brand.get('title','HSE Assistant')}</div>
    <div class="app-sub">{brand.get('subtitle','')}</div>
  </div>
  <div class="app-badge">{n_chunks:,} chunks indexed</div>
</div>
""", unsafe_allow_html=True)

missing = [k for k, v in {"MISTRAL_API_KEY": MISTRAL_KEY, "QDRANT_URL": QDRANT_URL,
                          "QDRANT_API_KEY": QDRANT_KEY}.items() if not v]
if missing:
    st.error("Missing configuration: " + ", ".join(missing) + ". Set them in the app's Secrets.")
    st.stop()

tab_chat, tab_admin = st.tabs(["💬  Ask", "⚙️  Admin"])

# ── CHAT ──
with tab_chat:
    st.session_state.setdefault("messages", [])
    if n_chunks == 0:
        st.info("The knowledge base is empty. An admin needs to run **Sync** (or ingest.py) first.")

    for m in st.session_state.messages:
        with st.chat_message(m["role"], avatar=("🧑" if m["role"] == "user" else "🦺")):
            st.markdown(m["content"])

    if not st.session_state.messages and n_chunks:
        st.markdown('<div class="suggest-label">Try asking</div>', unsafe_allow_html=True)
        cols = st.columns(2)
        for i, s in enumerate(SUGGESTED):
            if cols[i % 2].button(s, key=f"sug{i}", use_container_width=True):
                st.session_state._pending = s
                st.rerun()

    q = st.chat_input("Ask about HSE procedures, standards, PPE, permits…")
    if not q and st.session_state.get("_pending"):
        q = st.session_state.pop("_pending")
    if q:
        st.session_state.messages.append({"role": "user", "content": q})
        with st.chat_message("user", avatar="🧑"):
            st.markdown(q)
        with st.chat_message("assistant", avatar="🦺"):
            with st.spinner("Searching documents…"):
                try:
                    text, refs = answer(q)
                except Exception as e:
                    text, refs = f"Error: {e}", []
            if refs:
                text += "\n\n**Sources**\n" + "\n".join(
                    f"- [{r['name']}]({r['web_url']})" if r["web_url"] else f"- {r['name']}"
                    for r in refs)
            st.markdown(text)
            st.session_state.messages.append({"role": "assistant", "content": text})

# ── ADMIN ──
with tab_admin:
    if ADMIN_PW and not st.session_state.get("is_admin"):
        st.markdown("#### Admin sign-in")
        pw = st.text_input("Admin password", type="password")
        if st.button("Unlock"):
            st.session_state.is_admin = (pw == ADMIN_PW); st.rerun()
        st.stop()

    # Branding
    st.markdown("#### 🎨 Branding")
    c1, c2 = st.columns([1, 2])
    with c1:
        if src:
            st.image(src, width=120)
        else:
            st.caption("No logo yet")
    with c2:
        up = st.file_uploader("Upload company logo (PNG/JPG/SVG, keep it small)",
                              type=["png", "jpg", "jpeg", "svg"])
        title = st.text_input("App / company name", value=brand.get("title", "HSE Assistant"))
        subtitle = st.text_input("Subtitle", value=brand.get("subtitle", ""))
        accent = st.color_picker("Accent colour", value=brand.get("accent", "#0F9D8C"))
    b1, b2 = st.columns(2)
    if b1.button("💾 Save branding", use_container_width=True):
        new = dict(brand); new.update({"title": title, "subtitle": subtitle, "accent": accent})
        if up is not None:
            data = up.read()
            if len(data) > 1_500_000:
                st.warning("Logo is large (>1.5 MB). Please upload a smaller image.")
            else:
                new["logo_b64"] = base64.b64encode(data).decode()
                new["logo_mime"] = up.type or "image/png"
        try:
            save_brand(new); st.success("Saved. Refresh to see the new branding.")
        except Exception as e:
            st.error(f"Could not save: {e}")
    if b2.button("Remove logo", use_container_width=True):
        new = dict(brand); new["logo_b64"] = ""; new["logo_mime"] = ""; new["logo_url"] = ""
        save_brand(new); st.success("Logo removed. Refresh to update.")

    st.divider()
    # Diagnostics
    st.markdown("#### 🩺 Diagnostics")
    st.write("**Qdrant URL in use:**", f"`{QDRANT_URL}`")
    if st.button("🧪 Test Qdrant connection"):
        try:
            cols = [c.name for c in qdrant().get_collections().collections]
            st.success(f"Qdrant OK. Collections: {cols or '(none yet)'}")
        except Exception as e:
            st.error(f"Qdrant failed: {e}")

    st.divider()
    # Sync (memory-safe, slice per click)
    st.markdown("#### 🔄 Sync documents (memory-safe)")
    st.caption("For the full 1900-doc load, run **ingest.py** (locally or via GitHub Actions). "
               "In-app Sync indexes a slice per click and resumes (never restarts from doc 1).")
    st.write(f"**Site:** {SITE_URL or '(SITE_URL not set)'}")
    st.write(f"**Index:** {index_count():,} chunks")
    batch_n = st.number_input("Files to index per click", 20, 300, 100, step=20)
    force = st.checkbox("Force re-index (re-chunk even already-indexed files)")
    if not (TENANT_ID and CLIENT_ID and CLIENT_SECRET and SITE_URL):
        st.warning("Set TENANT_ID, CLIENT_ID, CLIENT_SECRET and SITE_URL in Secrets to enable Sync.")
    else:
        if st.button("🔄 Index next batch", use_container_width=True):
            bar = st.progress(0.0); status = st.empty(); logbox = st.empty(); buf = []
            def log(m):
                buf.append(str(m)); logbox.code("\n".join(buf[-25:]))
            def progress(done, total, indexed, notext, failed, chunks):
                bar.progress(done / total if total else 1.0)
                status.markdown(f"**{done}/{total}** this batch · {indexed} indexed · "
                                f"{notext} no-text · {failed} failed · {chunks} chunks")
            try:
                remaining = sync_batch(int(batch_n), progress, log, force=force)
                if remaining == 0:
                    st.success("All documents indexed. 🎉")
                else:
                    st.info(f"Batch done. ~{remaining} files still to index — click **Index next batch** again.")
            except Exception as e:
                log(f"ERROR: {e}")

    st.divider()
    # Danger zone
    st.markdown("#### 🗑️ Danger zone")
    st.caption("Clear removes ALL vectors so the next Sync/ingest re-chunks every document.")
    confirm = st.checkbox("I understand — wipe the whole index")
    if st.button("Clear entire index", disabled=not confirm):
        try:
            qdrant().delete_collection(COLLECTION); ensure_collection()
            st.success("Index cleared. Run Sync or ingest.py to rebuild.")
        except Exception as e:
            st.error(str(e))
